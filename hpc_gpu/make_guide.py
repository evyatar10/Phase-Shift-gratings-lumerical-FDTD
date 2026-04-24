"""Run this script once to generate the Word document guide."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

BLUE_DARK  = RGBColor(0x1F, 0x49, 0x7D)
BLUE_MID   = RGBColor(0x2E, 0x74, 0xB5)
ORANGE     = RGBColor(0xC0, 0x50, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG    = 'F0F0F0'
HDR_BG     = '1F497D'

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BLUE_DARK
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE_MID
    return p

def code(text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line.strip() else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  CODE_BG)
        p._p.get_or_add_pPr().append(shd)

def body(text):
    doc.add_paragraph(text)

def bullet(text, bold=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold:
        r = p.add_run(bold)
        r.bold = True
    p.add_run(text)

def note(text):
    p = doc.add_paragraph()
    r = p.add_run('NOTE: ')
    r.bold = True
    r.font.color.rgb = ORANGE
    p.add_run(text)

def sp():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)

def tbl(rows):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, val in enumerate(rows[0]):
        hdr[i].text = val
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = WHITE
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  HDR_BG)
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shd)
    for row_data in rows[1:]:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

# ── Title ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Athena GPU Workflow')
r.font.name = 'Calibri'; r.font.size = Pt(28); r.bold = True; r.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Setup & Usage Guide').font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Technion HPC  |  Lumerical FDTD 2026 R1.1  |  NVIDIA A100').font.size = Pt(12)

doc.add_paragraph()

# ── Prerequisites ──────────────────────────────────────────────────────────────
h1('Prerequisites  (Already Completed)')
tbl([
    ['Item', 'Status'],
    ['Athena account', 'Yes'],
    ['License server reachable from Athena', 'Yes  (1055@132.68.48.51)'],
    ['Local Lumerical version', '2026 R1.1  (v261)'],
    ['Workflow code in repo', 'hpc_gpu/'],
])

# ── Step 0 ─────────────────────────────────────────────────────────────────────
h1('Step 0 — One-Time Cluster Checks')
body('SSH into Athena:')
code('ssh evyatarrubin@dgx-master.technion.ac.il')
sp()

h2('0a.  Check Apptainer is available')
body('Apptainer is needed to build the container image.')
code('apptainer --version')
body('If missing: build the container on WSL2 instead (see Step 1).')
sp()

h2('0b.  Find your scratch / storage path')
code('df -h\nls /scratch/$USER 2>/dev/null || ls /raid/$USER 2>/dev/null')
body('If a large scratch path exists, update these two values:')
bullet('hpc_gpu/deploy_athena.sh  →  REMOTE_BASE variable')
bullet('hpc_gpu/scripts/athena_run.py  →  BASE_SAVE_DIR variable')
body('Otherwise results go to ~/bragg_sim_gpu/  (home, usually ~100 GB quota).')
sp()

h2('0c.  Check if a billing account code is required')
code('sbatch --wrap="echo hello"')
body('If the command fails mentioning "--account", ask Technion CIS for your account code and add this line to all three files in hpc_gpu/jobs/:')
code('#SBATCH --account=<your_code>')
body('Files: run_fsp_gpu.sh  |  run_fsp_gpu_array.sh  |  run_python_gpu.sh')
sp()

# ── Step 1 ─────────────────────────────────────────────────────────────────────
h1('Step 1 — Verify Lumerical Linux Install (Already Done)')
bullet('Lumerical 2026 R1.1 was installed in WSL2 via AnsysInstaller.sh')
bullet('Install location: ~/ansys_incS/v261/Lumerical/')
bullet('The container build copies these files in — no re-download or auth needed')
sp()
body('To verify:')
code('ls ~/ansys_incS/v261/Lumerical/bin/fdtd-engine-ompi-lcl')
sp()

# ── Step 2 ─────────────────────────────────────────────────────────────────────
h1('Step 2 — Build and Upload the Container')
body('Done ONCE from WSL2. The build copies the installed Lumerical files into the image.')
sp()
h2('If Apptainer is not yet installed in WSL2:')
code('sudo add-apt-repository ppa:apptainer/ppa\nsudo apt install -y apptainer')
sp()
h2('Then build and upload:')
code('cd hpc_gpu/container\nbash build.sh')
sp()
body('This script will automatically:')
bullet('Build lumerical-2026R1.sif  (5–15 minutes — file copy, no download)')
bullet('Convert to lumerical-2026R1.sqsh  (Athena\'s Enroot format)')
bullet('Upload the .sqsh to ~/containers/ on Athena')
sp()
note('You only need to repeat this if you upgrade Lumerical versions.')
sp()

# ── Step 3 ─────────────────────────────────────────────────────────────────────
h1('Step 3 — Smoke Test the Container')
body('On the Athena login node, verify the engine binary works:')
code('srun --gpus=1 \\\n  --container-image=$HOME/containers/lumerical-2026R1.sqsh \\\n  --container-env=ANSYSLMD_LICENSE_FILE=1055@132.68.48.51 \\\n  /opt/lumerical/v261/bin/fdtd-engine-ompi-lcl -v')
body('Expected: Lumerical version string printed, no errors.')
sp()
body('Also check your license seat count and tier:')
code('srun --gpus=1 \\\n  --container-image=$HOME/containers/lumerical-2026R1.sqsh \\\n  --container-env=ANSYSLMD_LICENSE_FILE=1055@132.68.48.51 \\\n  /opt/lumerical/v261/bin/lmutil lmstat -a -c 1055@132.68.48.51')
body('Look for "fdtd_gpu" in the output. The number next to it = how many simulations can run simultaneously. Update K=4 in hpc_gpu/deploy_athena.sh to match this number.')
sp()

# ── Step 4 ─────────────────────────────────────────────────────────────────────
h1('Step 4 — First Real Run  (Single Simulation, Engine-Only)')
body('Run from your LOCAL Windows machine in Git Bash or WSL:')
code('bash hpc_gpu/deploy_athena.sh --option1 --preset single --watch')
sp()
body('This will:')
bullet('Generate the .fsp file locally  (same as Zeus)')
bullet('Upload it to Athena')
bullet('Submit the SLURM GPU job')
bullet('Poll every 60 seconds until done')
bullet('Download results to:  results_from_athena/')
sp()

# ── Step 5 ─────────────────────────────────────────────────────────────────────
h1('Step 5 — Benchmark')
body('Run the SAME simulation on both clusters and compare wall times:')
sp()
tbl([
    ['', 'Zeus (CPU)', 'Athena (GPU)'],
    ['Command', 'bash hpc/deploy.sh --option1 --preset single', 'bash hpc_gpu/deploy_athena.sh --option1 --preset single'],
    ['Scheduler', 'PBS', 'SLURM'],
    ['Hardware', '80 CPU cores', '1× A100 GPU'],
    ['Cost', 'Free', 'Billed per GPU-hour'],
])
body('Reference benchmark (Ansys metalens, 0.85 billion Yee cells):')
tbl([
    ['Hardware', 'Wall Time', 'Speedup'],
    ['63-core CPU  (≈ 1 Zeus node)', '62 minutes', '1.0×  (baseline)'],
    ['1× A100 GPU', '12 minutes', '5.2×'],
    ['4× A100 GPU', '9 minutes', '7.1×'],
    ['8× A100 GPU', '6.6 minutes', '9.4×'],
])
note('If your grating simulation is small (<50 million cells), GPU speedup may be minimal. Keep small jobs on Zeus — it is free.')
sp()

# ── Step 6 ─────────────────────────────────────────────────────────────────────
h1('Step 6 — Parameter Sweep  (Job Array)')
code('bash hpc_gpu/deploy_athena.sh --option1 --preset sweep_shift --watch\nbash hpc_gpu/deploy_athena.sh --option1 --preset sweep_inner_size --watch')
sp()
body('When multiple .fsp files are detected, deploy_athena.sh automatically submits them as a SLURM job array — all running in parallel on separate GPUs, throttled by your license seat count.')
sp()
note('This is the single biggest upgrade over Zeus, which runs sweep files serially inside one job.')
sp()

# ── Step 7 ─────────────────────────────────────────────────────────────────────
h1('Step 7 — Full Python Pipeline  (Phase 2, Optional)')
body('Only attempt this after Step 4 is working reliably.')
code('bash hpc_gpu/deploy_athena.sh --option2 --run single_sim --watch\nbash hpc_gpu/deploy_athena.sh --option2 --run sweep_shift --watch\nbash hpc_gpu/deploy_athena.sh --option2 --run sweep_inner_size --watch')
sp()
body('Runs the full lumapi Python pipeline inside the container with USE_GPU=True. GPU is enabled automatically on every FDTD session via athena_run.py — no changes needed to your simulation modules.')
sp()

# ── Daily usage ────────────────────────────────────────────────────────────────
h1('Daily Usage  (After Setup is Complete)')
tbl([
    ['Task', 'Command'],
    ['Single sim + wait for results', 'bash hpc_gpu/deploy_athena.sh --option1 --preset single --watch'],
    ['Sweep + wait for results', 'bash hpc_gpu/deploy_athena.sh --option1 --preset sweep_shift --watch'],
    ['Download results only', 'bash hpc_gpu/deploy_athena.sh --results'],
    ['Check running jobs', 'bash hpc_gpu/deploy_athena.sh --watch-only'],
])

# ── Monitoring ─────────────────────────────────────────────────────────────────
h1('Monitoring on Athena Directly')
body('SSH in:')
code('ssh evyatarrubin@dgx-master.technion.ac.il')
sp()
tbl([
    ['Command', 'Purpose'],
    ['squeue -u evyatarrubin', 'View all your running / pending jobs'],
    ['squeue -j <job_id> -l', 'Details on one specific job'],
    ['tail -f ~/bragg_sim_gpu/jobs/logs/*.out', 'Live log output of latest job'],
    ['scancel <job_id>', 'Cancel a job'],
])

# ── Scaling ────────────────────────────────────────────────────────────────────
h1('Scaling GPUs')
body('Edit hpc_gpu/jobs/run_fsp_gpu.sh — change #SBATCH --gpus=N and NUM_MPI_RANKS=N to match:')
tbl([
    ['Configuration', 'Expected Speedup', 'When to Use'],
    ['--gpus=1  (1× A100)', '≈5× vs 1 Zeus node', 'Start here — benchmark first'],
    ['--gpus=4  (4× A100)', '≈7×', 'Sims larger than ~200M cells'],
    ['--gpus=8  (full DGX node)', '≈9×', 'Very large sims'],
    ['--gpus=16+  (multi-node)', 'Requires Enterprise license', 'Verify with lmutil lmstat first'],
])

# ── Key files ──────────────────────────────────────────────────────────────────
h1('Key Files')
tbl([
    ['File', 'Purpose'],
    ['hpc_gpu/deploy_athena.sh', 'Run locally to submit jobs to Athena'],
    ['hpc_gpu/container/build.sh', 'Run once to build and upload the container'],
    ['hpc_gpu/container/lumerical.def', 'Container recipe (shareable with colleagues)'],
    ['hpc_gpu/jobs/run_fsp_gpu.sh', 'SLURM script — single GPU run'],
    ['hpc_gpu/jobs/run_fsp_gpu_array.sh', 'SLURM script — sweep job array'],
    ['hpc_gpu/jobs/run_python_gpu.sh', 'SLURM script — full Python pipeline'],
    ['hpc_gpu/scripts/athena_run.py', 'Server-side Python dispatcher (GPU-aware)'],
    ['results_from_athena/', 'Local folder where results are downloaded'],
])

# ── Zeus untouched ─────────────────────────────────────────────────────────────
h1('Zeus Workflow — Not Affected')
body('Everything in hpc/ is completely unchanged. Zeus runs continue exactly as before:')
code('bash hpc/deploy.sh --option1    # Zeus, CPU, unaffected\nbash hpc/deploy.sh --option2    # Zeus, CPU, unaffected')
sp()

# ── Troubleshooting ────────────────────────────────────────────────────────────
h1('Troubleshooting')
tbl([
    ['Symptom', 'Fix'],
    ['"Cannot connect to display"', 'Rebuild the container: bash hpc_gpu/container/build.sh'],
    ['"License server not responding"', 'Run: nc -vz 132.68.48.51 1055  from the compute node'],
    ['"fdtd-engine-ompi-lcl: not found"', 'Container did not install correctly — check build logs and rebuild'],
    ['"Out of GPU memory"', 'Reduce mesh size, or increase --gpus to spread across more A100s'],
    ['"setresource GPU error"', 'License may not include fdtd_gpu feature — check lmutil lmstat'],
    ['"sbatch: Batch job submission failed"', 'Add  #SBATCH --account=<code>  to job scripts (ask CIS)'],
    ['"Job pending for a long time"', 'Athena has 72 A100s total (9 nodes x 8). If all busy, you wait in queue.'],
])

out = r'C:\Users\evyat\Documents\Athena_GPU_Setup_Guide.docx'
doc.save(out)
print('Saved:', out)
